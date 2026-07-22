"""评审报告 Word 导出。"""

from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.schemas.review_report import ReviewReportDto


def render_review_report_docx(report: ReviewReportDto) -> bytes:
    doc = Document()
    _configure_default_style(doc)

    title = doc.add_heading(report.title, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. 基本信息", level=1)
    _add_basic_info_table(doc, report)

    doc.add_heading("2. 缺失特性结论", level=1)
    _add_missing_features_table(doc, report)

    doc.add_heading("3. 需提升结论", level=1)
    _add_improvements_table(doc, report)

    doc.add_heading("4. 评审总结建议", level=1)
    _add_summary_points(doc, report)

    doc.add_heading("5. 最终结论", level=1)
    doc.add_paragraph(report.finalConclusion or "（无）")

    if report.citations:
        doc.add_heading("引用依据", level=1)
        for cite in report.citations[:20]:
            line = " ".join(
                part
                for part in [
                    cite.standardNo,
                    cite.standardName,
                    cite.position,
                    cite.excerpt[:200] if cite.excerpt else "",
                ]
                if part
            )
            doc.add_paragraph(line or "—", style="List Bullet")

    doc.add_paragraph("")
    disclaimer = doc.add_paragraph(report.disclaimer)
    disclaimer.runs[0].italic = True

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _configure_default_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(11)


def _add_basic_info_table(doc: Document, report: ReviewReportDto) -> None:
    info = report.basicInfo
    rows = [
        ("标准名称", info.standardName or "—"),
        ("标准编号", info.standardNo or "—"),
        ("所属零部件", info.partName or "—"),
        ("评审日期", info.reviewDate or "—"),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    for index, (label, value) in enumerate(rows):
        table.rows[index].cells[0].text = label
        table.rows[index].cells[1].text = value


def _add_missing_features_table(doc: Document, report: ReviewReportDto) -> None:
    headers = ["缺失特性说明", "涉及标准"]
    body_rows = [
        [row.description, row.relatedStandards or "—"]
        for row in report.missingFeatures
    ]
    _add_data_table(doc, headers, body_rows, empty_text="未识别到缺失特性")


def _add_improvements_table(doc: Document, report: ReviewReportDto) -> None:
    headers = ["特性要素", "建议内容", "建议方法", "参考标准"]
    body_rows = [
        [
            row.featureElement,
            row.suggestedContent or "—",
            row.suggestedMethod or "—",
            row.referenceStandards or "—",
        ]
        for row in report.improvements
    ]
    _add_data_table(doc, headers, body_rows, empty_text="未识别到需提升项")


def _add_data_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    empty_text: str,
) -> None:
    if not rows:
        doc.add_paragraph(empty_text)
        return

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    for col_index, header in enumerate(headers):
        table.rows[0].cells[col_index].text = header

    for row_index, row_values in enumerate(rows, start=1):
        for col_index, value in enumerate(row_values):
            table.rows[row_index].cells[col_index].text = value


def _add_summary_points(doc: Document, report: ReviewReportDto) -> None:
    if not report.summaryPoints:
        doc.add_paragraph("暂无总结建议")
        return

    for point in report.summaryPoints:
        if point.title and point.content:
            text = f"{point.order}. {point.title}：{point.content}"
        elif point.title:
            text = f"{point.order}. {point.title}"
        else:
            text = f"{point.order}. {point.content}"
        doc.add_paragraph(text)
