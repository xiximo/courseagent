import json
import zipfile
from io import BytesIO

from docx import Document

from app.processing.docx_converter import convert_docx_to_markdown
from app.processing.docx_payload import prepare_docx_payload


def _make_docx() -> bytes:
    doc = Document()
    doc.add_heading("测试标题", level=1)
    doc.add_paragraph("正文内容")
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_convert_normal_docx() -> None:
    result = convert_docx_to_markdown(_make_docx(), "test.docx")
    assert result.char_count > 0


def test_reject_json_payload() -> None:
    payload = json.dumps({"ok": False, "message": "文件不存在"}).encode()
    try:
        prepare_docx_payload(payload, "bad.docx")
        raise AssertionError("expected ValueError")
    except ValueError as exc:
        assert "JSON" in str(exc)


def test_unwrap_zip_wrapper() -> None:
    good = _make_docx()
    outer = BytesIO()
    with zipfile.ZipFile(outer, "w") as zf:
        zf.writestr("inner/test.docx", good)
    prepared = prepare_docx_payload(outer.getvalue(), "test.docx")
    with zipfile.ZipFile(BytesIO(prepared), "r") as zf:
        assert "[Content_Types].xml" in zf.namelist()


def test_repair_null_relationship() -> None:
    good = _make_docx()
    broken_buf = BytesIO()
    with zipfile.ZipFile(BytesIO(good), "r") as zin:
        with zipfile.ZipFile(broken_buf, "w") as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "word/_rels/document.xml.rels":
                    text = data.decode("utf-8")
                    broken_rel = (
                        '<Relationship Id="rId99" '
                        'Type="http://schemas.openxmlformats.org/'
                        'officeDocument/2006/relationships/image" '
                        'Target="NULL"/>'
                    )
                    text = text.replace("</Relationships>", f"{broken_rel}</Relationships>")
                    data = text.encode("utf-8")
                zout.writestr(item, data)

    result = convert_docx_to_markdown(broken_buf.getvalue(), "broken.docx")
    assert result.char_count > 0
